"""
EIA Toolkit — AWS Lambda Report Generator
Generates government-compliant Word/PDF biodiversity chapters.

Deploy:
  pip install python-docx boto3 -t ./package
  cp lambda_function.py ./package/
  cd package && zip -r ../eia-report.zip .
  aws lambda create-function \
    --function-name eia-report-generator \
    --runtime python3.12 \
    --handler lambda_function.lambda_handler \
    --zip-file fileb://eia-report.zip \
    --role arn:aws:iam::YOUR_ACCOUNT:role/lambda-eia-role \
    --region ap-northeast-1 \
    --memory-size 512 \
    --timeout 60
  aws lambda add-permission \
    --function-name eia-report-generator \
    --statement-id public-url \
    --action lambda:InvokeFunctionUrl \
    --principal "*" \
    --function-url-auth-type NONE
  aws lambda create-function-url-config \
    --function-name eia-report-generator \
    --auth-type NONE
"""

import json
import base64
import io
import os
from datetime import datetime, date

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ── Red List categories (Japanese) ───────────────────────────────────────────
RL_LABELS = {
    "CR": "絶滅危惧IA類（CR）",
    "EN": "絶滅危惧IB類（EN）",
    "VU": "絶滅危惧II類（VU）",
    "NT": "準絶滅危惧（NT）",
    "LC": "軽度懸念（LC）",
    "EX": "絶滅（EX）",
}

RL_PROTECTED_LAWS = {
    "CR": "種の保存法・文化財保護法等により保護対象となる場合がある。",
    "EN": "種の保存法による国内希少野生動植物種に指定されている場合がある。",
    "VU": "鳥獣保護管理法・各都道府県条例等により規制対象となる場合がある。",
    "NT": "引き続きモニタリングが必要な種。",
    "LC": "現時点では絶滅危惧に該当しない。",
}

# ── Stage labels ──────────────────────────────────────────────────────────────
STAGE_LABELS = {
    1: "配慮書手続",
    2: "方法書手続",
    3: "現地調査",
    4: "準備書手続",
    5: "意見聴取",
    6: "評価書",
    7: "事後調査",
}


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1B, 0x43, 0x32)
    return p


def add_para(doc, text: str, indent: bool = False):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


# ── Core report builder ───────────────────────────────────────────────────────

def build_report(project: dict, species_list: list, report_type: str) -> bytes:
    """
    Build a Word document and return it as bytes.
    report_type: 'preparatory' | 'assessment' | 'tnfd'
    """
    doc = Document()

    # Page setup — A4
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    pref    = project.get("pref", "")
    name    = project.get("name", "（事業名未設定）")
    client  = project.get("client", "")
    area    = project.get("area", "")
    stage   = project.get("stage", 4)
    manager = project.get("manager", "")
    today   = datetime.now().strftime("%Y年%m月%d日")

    chapter_num = "第4章" if report_type == "preparatory" else "第5章"

    # ── Title block ──────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"{chapter_num}　生物・生態系への影響")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x43, 0x32)

    doc.add_paragraph()  # spacer

    # Project info table
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ("事業名",      name),
        ("事業者",      client),
        ("対象地域",    f"{pref}　（事業面積：{area} ha）"),
        ("報告書作成日", today),
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        set_cell_bg(info_table.rows[i].cells[0], "D8EFE3")
        for cell in info_table.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # ── 4.1 調査の概要 ────────────────────────────────────────────────────────
    add_heading(doc, f"{chapter_num[:-1]}.1　現地調査の概要", level=2)
    add_para(doc,
        f"本章は、環境影響評価法（平成9年法律第81号）第{stage}段階手続に係る"
        f"生物・生態系分野の調査結果及び環境影響評価の結果を取りまとめたものである。"
        f"現地調査は{pref}内の事業区域（{area} ha）において実施した。"
    )
    add_para(doc,
        "調査は環境省「自然環境保全基礎調査」及び「環境影響評価技術指針（動植物・生態系）」"
        "（平成28年版）に準拠した手法により実施した。"
    )

    # ── 4.2 確認種の概要 ──────────────────────────────────────────────────────
    add_heading(doc, f"{chapter_num[:-1]}.2　確認種の概要", level=2)

    total = len(species_list)
    rl_species = [s for s in species_list if s.get("status") in ("CR","EN","VU","NT")]
    protected  = [s for s in species_list if s.get("protected")]

    # Summary stats paragraph
    add_para(doc,
        f"現地調査の結果、合計{total}種の動植物が確認された。"
        f"このうち、環境省レッドリスト（2020年版）掲載種は{len(rl_species)}種、"
        f"種の保存法等により保護指定を受けている種は{len(protected)}種であった。"
    )

    # Category breakdown table
    doc.add_paragraph()
    p = doc.add_paragraph("■ レッドリストカテゴリ別確認種数")
    p.runs[0].font.bold = True

    cat_table = doc.add_table(rows=1, cols=3)
    cat_table.style = 'Table Grid'
    headers = ["カテゴリ", "カテゴリ名", "確認種数"]
    for i, h in enumerate(headers):
        cat_table.rows[0].cells[i].text = h
        set_cell_bg(cat_table.rows[0].cells[i], "1B4332")
        for run in cat_table.rows[0].cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold = True
            run.font.size = Pt(9)

    for cat in ["CR","EN","VU","NT","LC"]:
        count = len([s for s in species_list if s.get("status") == cat])
        row = cat_table.add_row()
        row.cells[0].text = cat
        row.cells[1].text = RL_LABELS.get(cat, cat)
        row.cells[2].text = f"{count}種"
        if cat in ("CR","EN"):
            set_cell_bg(row.cells[0], "FEE2E2")
        elif cat == "VU":
            set_cell_bg(row.cells[0], "FEF3C7")
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)

    doc.add_paragraph()

    # ── 4.3 確認種一覧 ────────────────────────────────────────────────────────
    add_heading(doc, f"{chapter_num[:-1]}.3　確認種一覧", level=2)
    add_para(doc, "現地調査により確認された全種の一覧を以下に示す。")

    if species_list:
        sp_table = doc.add_table(rows=1, cols=7)
        sp_table.style = 'Table Grid'
        sp_headers = ["No.", "種名（和名）", "学名", "分類群",
                      "環境省RL", "保護指定", "確認個体数"]
        for i, h in enumerate(sp_headers):
            sp_table.rows[0].cells[i].text = h
            set_cell_bg(sp_table.rows[0].cells[i], "2D6A4F")
            for run in sp_table.rows[0].cells[i].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(8.5)

        for idx, sp in enumerate(species_list, 1):
            row = sp_table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = sp.get("name", "")
            row.cells[2].text = sp.get("latin", "")
            row.cells[3].text = sp.get("type", "")
            row.cells[4].text = sp.get("status", "LC")
            row.cells[5].text = "○" if sp.get("protected") else "−"
            row.cells[6].text = str(sp.get("count", ""))

            # Highlight protected species rows
            if sp.get("protected") or sp.get("status") in ("CR","EN"):
                for cell in row.cells:
                    set_cell_bg(cell, "FFF5F5")

            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8.5)
    else:
        add_para(doc, "（確認種データが未入力です。第3段階現地調査後に再生成してください。）",
                 indent=True)

    doc.add_paragraph()

    # ── 4.4 重要種の詳細 ──────────────────────────────────────────────────────
    add_heading(doc, f"{chapter_num[:-1]}.4　重要種の生息状況及び影響予測", level=2)

    if rl_species:
        for sp in rl_species:
            add_heading(doc, f"（{sp.get('name','不明')}）", level=3)
            cat = sp.get("status","LC")
            add_para(doc, f"【保全上の位置づけ】環境省レッドリスト {RL_LABELS.get(cat,cat)}")
            add_para(doc, f"【確認状況】{sp.get('location','確認地点未記入')}にて"
                         f"{sp.get('count',1)}個体を確認。確認日：{sp.get('obs_date','未記録')}")
            if sp.get("notes"):
                add_para(doc, f"【生態・行動メモ】{sp['notes']}", indent=True)
            add_para(doc,
                f"【影響予測】工事による生息地の直接損失及び騒音・振動の影響が懸念される。"
                f"事業者は環境保全措置として、工事時期の制限、緩衝緑地の設置等を検討する必要がある。"
                f"{RL_PROTECTED_LAWS.get(cat,'')}"
            )
            add_para(doc,
                "【環境保全措置】工事施工前に生息確認調査を再実施し、個体が確認された場合は"
                "専門家の指導のもと適切な保護措置（移植・迂回等）を講じること。"
            )
    else:
        add_para(doc,
            "現地調査において、環境省レッドリスト掲載種は確認されなかった。"
            "引き続き工事中モニタリングにより動植物相の変化を監視する。"
        )

    doc.add_paragraph()

    # ── 4.5 生態系への総合評価 ────────────────────────────────────────────────
    add_heading(doc, f"{chapter_num[:-1]}.5　生態系への総合評価", level=2)
    add_para(doc,
        f"事業区域（{pref}、{area} ha）は、"
        f"{'重要な動植物種の生息地を含み、慎重な環境配慮が求められる地域である。' if rl_species else '生物多様性の観点から比較的低リスクな地域と評価される。'}"
    )
    add_para(doc,
        f"確認種{total}種のうちレッドリスト掲載種{len(rl_species)}種について、"
        "個別の環境保全措置を実施することにより、事業による生物多様性への影響を"
        "最小限に抑えることが可能と判断される。"
    )
    add_para(doc,
        "事後調査計画に基づき、工事中及び供用後においても継続的なモニタリングを実施し、"
        "予期せぬ影響が生じた場合は速やかに追加的保全措置を検討する。"
    )

    # ── TNFD additional section ───────────────────────────────────────────────
    if report_type == "tnfd":
        doc.add_page_break()
        add_heading(doc, "【TNFD LEAP整合】自然関連財務情報開示", level=1)
        add_para(doc,
            "本節は、自然関連財務情報開示タスクフォース（TNFD）のLEAPアプローチ"
            "（Locate・Evaluate・Assess・Prepare）に基づく情報整理である。"
        )
        for step, content in [
            ("Locate（所在地特定）",
             f"事業区域：{pref}（面積：{area} ha）。"
             "自然環境への接触・依存度：高（生物多様性豊かな地域に立地）。"),
            ("Evaluate（評価）",
             f"確認種{total}種のうちRL掲載種{len(rl_species)}種。"
             "生態系サービス：水資源涵養・炭素固定機能への依存度が高い。"),
            ("Assess（影響・依存関係の評価）",
             "主要なネイチャーリスク：生息地損失（直接的）・水質変化（間接的）。"
             "財務的影響：許認可遅延リスク、訴訟リスク、レピュテーションリスク。"),
            ("Prepare（対応策の策定）",
             "保全措置実施、TNFD整合開示、自然ポジティブ目標の設定を推奨する。"),
        ]:
            add_heading(doc, step, level=2)
            add_para(doc, content)

    # ── Footer note ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    note = doc.add_paragraph(
        f"※本報告書は EIAツールキット（ReCorta LLC）により自動生成されました。"
        f"内容は専門家による最終確認を経てから提出してください。"
        f"生成日時：{today}　担当者：{manager}"
    )
    for run in note.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x78, 0x71, 0x6C)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Lambda handler ────────────────────────────────────────────────────────────

def lambda_handler(event, context):
    # CORS preflight
    if event.get("requestContext", {}).get("http", {}).get("method") == "OPTIONS":
        return {
            "statusCode": 200,
            "headers": cors_headers(),
            "body": "",
        }

    try:
        body = json.loads(event.get("body", "{}"))
    except (json.JSONDecodeError, TypeError):
        return error_response(400, "Invalid JSON body")

    project      = body.get("project", {})
    species_list = body.get("species", [])
    report_type  = body.get("report_type", "preparatory")  # preparatory | assessment | tnfd

    if not project:
        return error_response(400, "Missing 'project' field in body")

    if not DOCX_AVAILABLE:
        return error_response(500, "python-docx not installed in Lambda package")

    try:
        docx_bytes = build_report(project, species_list, report_type)
    except Exception as e:
        return error_response(500, f"Report generation failed: {str(e)}")

    # Return base64-encoded Word document
    encoded = base64.b64encode(docx_bytes).decode("utf-8")
    stage   = project.get("stage", 4)
    pref    = project.get("pref", "")
    fname   = (
        f"評価書_生物多様性章_{pref}.docx" if report_type == "assessment"
        else f"準備書_生物多様性章_{pref}.docx" if report_type == "preparatory"
        else f"TNFD_LEAP整合_{pref}.docx"
    )

    return {
        "statusCode": 200,
        "headers": {
            **cors_headers(),
            "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "Content-Disposition": f'attachment; filename="{fname}"',
        },
        "body": encoded,
        "isBase64Encoded": True,
    }


def cors_headers():
    return {
        "Access-Control-Allow-Origin": "*",   # Restrict to your Vercel domain in production
        "Access-Control-Allow-Methods": "POST, OPTIONS",
        "Access-Control-Allow-Headers": "Content-Type",
    }


def error_response(status: int, message: str):
    return {
        "statusCode": status,
        "headers": cors_headers(),
        "body": json.dumps({"error": message}),
    }


# ── Local test ────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    test_event = {
        "requestContext": {"http": {"method": "POST"}},
        "body": json.dumps({
            "report_type": "preparatory",
            "project": {
                "name": "北海道洋上風力発電EIA",
                "client": "J-Power株式会社",
                "pref": "北海道",
                "area": "2400",
                "stage": 4,
                "manager": "田中 誠一",
            },
            "species": [
                {"name": "オジロワシ",   "latin": "Haliaeetus albicilla",
                 "type": "鳥類", "status": "VU", "protected": True,  "count": 2,
                 "location": "調査地点A-3", "obs_date": "2026-05-12",
                 "notes": "営巣確認。繁殖期中の工事は避けること。"},
                {"name": "エゾシカ",    "latin": "Cervus nippon yesoensis",
                 "type": "哺乳類", "status": "LC", "protected": False, "count": 8,
                 "location": "調査地点B-1", "obs_date": "2026-05-14", "notes": ""},
                {"name": "タンチョウ",  "latin": "Grus japonensis",
                 "type": "鳥類", "status": "VU", "protected": True,  "count": 1,
                 "location": "調査地点A-7", "obs_date": "2026-05-18", "notes": "採餌行動を確認。"},
            ]
        })
    }
    result = lambda_handler(test_event, None)
    if result["isBase64Encoded"]:
        with open("test_output.docx", "wb") as f:
            f.write(base64.b64decode(result["body"]))
        print("✅ Written to test_output.docx")
    else:
        print("Error:", result["body"])
